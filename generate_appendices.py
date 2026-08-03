#!/usr/bin/env python3
"""
generate_appendices.py
======================
Generate the SPL traceability appendices (A / B / C) from a Test Specification
export and append them into a *copy* of that export.

  Appendix A - Requirement Traceability Matrix   (Requirement -> Test Case -> Test Suite)
  Appendix B - Test Suites Traceability Matrix    (Test Suite -> Description / Execution / Test Cases)
  Appendix C - Test Cases Traceability Matrix     (Test Case  -> Description / Test Suites)

The export is expected to be the standard QR_TS_SYS_* test-suite export (.doc or .docx),
where each suite is one table containing its description, Execution type, its
QR_TC_SYS_* test cases, and step rows that reference SYS_#### requirements.

The Safety-Related classification is NOT in the export, so Appendix A takes it
from an optional RVTM / SRS lookup file (--safety). Requirements not found in the
lookup (or when no lookup is given) are marked with the placeholder (default "TBC").

--------------------------------------------------------------------------------
USAGE
-----
  python generate_appendices.py EXPORT.docx
  python generate_appendices.py EXPORT.doc  --safety RVTM.xlsx
  python generate_appendices.py EXPORT.docx --safety SRS.xlsx --appendices ABC -o OUT.docx

Common options:
  --safety FILE        RVTM/SRS .xlsx or .csv giving each requirement's safety class
  --req-col NAME       Header name of the requirement-ID column (auto-detected if omitted)
  --safety-col NAME    Header name of the safety-class column   (auto-detected if omitted)
  --appendices ABC     Which appendices to build, any subset of A B C (default: ABC)
  --srs-rev "11.00 [5]" Text shown in Appendix A's safety column header
  --placeholder TBC    Value used when a requirement's safety class is unknown
  -o, --output FILE    Output path (default: <EXPORT>_with_appendices.docx)

Dependencies: python-docx  (+ LibreOffice `soffice` only if the input is a legacy .doc)
              openpyxl is used for .xlsx lookups (falls back gracefully if absent)
--------------------------------------------------------------------------------
"""

import argparse
import os
import re
import shutil
import subprocess
import sys
from collections import OrderedDict, defaultdict

try:
    import docx
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    sys.exit("ERROR: python-docx is required.  Install with:  pip install python-docx")

# --------------------------------------------------------------------------- #
#  Appendix definitions - edit here if column headings/titles ever change      #
# --------------------------------------------------------------------------- #
HEADER_FILL = "D9E2F3"      # light blue header shading
TC_RE  = re.compile(r'(QR_TC_SYS_\d+[A-Za-z]*)\s*-\s*(.*)', re.S)
TS_ID  = re.compile(r'QR_TS_SYS_[A-Za-z]+_\d+')
REQ_RE = re.compile(r'SYS_\d+')


def extract_reqs(text):
    """Requirement IDs, ignoring SYS_#### that is part of a QR_TC_SYS_/QR_TS_SYS_ id."""
    out = []
    for m in REQ_RE.finditer(text):
        pre = text[max(0, m.start() - 3):m.start()]
        if pre.endswith("TC_") or pre.endswith("TS_"):
            continue
        out.append(m.group(0))
    return out


# --------------------------------------------------------------------------- #
#  Input handling                                                              #
# --------------------------------------------------------------------------- #
def ensure_docx(path):
    """Return a .docx path; convert a legacy .doc via LibreOffice if needed."""
    if path.lower().endswith(".docx"):
        return path
    if not path.lower().endswith(".doc"):
        sys.exit(f"ERROR: unsupported input '{path}'. Provide a .doc or .docx file.")
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        sys.exit("ERROR: input is a legacy .doc and LibreOffice (soffice) was not found.\n"
                 "       Open it in Word and 'Save As' .docx, then re-run on the .docx.")
    outdir = os.path.dirname(os.path.abspath(path)) or "."
    print(f"  Converting legacy .doc -> .docx via LibreOffice ...")
    subprocess.run([soffice, "--headless", "--convert-to", "docx", "--outdir", outdir, path],
                   check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    converted = os.path.join(outdir, os.path.splitext(os.path.basename(path))[0] + ".docx")
    if not os.path.exists(converted):
        sys.exit("ERROR: .doc -> .docx conversion failed.")
    return converted


# --------------------------------------------------------------------------- #
#  Parsing the export                                                          #
# --------------------------------------------------------------------------- #
def _distinct(cells):
    out = []
    for c in cells:
        t = c.text.strip()
        if not out or out[-1] != t:
            out.append(t)
    return out


def parse_suite(tbl):
    """Return (description, execution, [ (tc_id, tc_desc, [req,...]), ... ])."""
    rows = tbl.rows
    desc = rows[1].cells[0].text.strip().replace("\n", " ") if len(rows) > 1 else ""
    desc = " ".join(desc.split())
    execu = ""
    for ri, row in enumerate(rows):
        texts = [c.text.strip() for c in row.cells]
        if "Execution" in texts and ri + 1 < len(rows):
            execu = rows[ri + 1].cells[texts.index("Execution")].text.strip()
            break
    tcs, cur = [], None
    for row in rows:
        first = row.cells[0].text.strip()
        if first == "TC":
            dd = _distinct(row.cells)
            m = TC_RE.match(dd[1]) if len(dd) > 1 else None
            if m:
                cur = [m.group(1), " ".join(m.group(2).split()), []]
                tcs.append(cur)
        elif cur is not None and re.match(r'^\d+', first):
            joined = " ".join(c.text for c in row.cells)
            for rid in extract_reqs(joined):
                if rid not in cur[2]:
                    cur[2].append(rid)
    return desc, execu, tcs


def parse_export(docx_path):
    """Return ordered list of (suite_id, description, execution, [tc,...])."""
    doc = docx.Document(docx_path)
    tables = doc.tables
    suites, cur_head, ti = [], None, 0
    for el in doc.element.body.iterchildren():
        tag = el.tag.split("}")[-1]
        if tag == "p":
            txt = "".join(t.text or "" for t in el.iter(qn("w:t"))).strip()
            m = TS_ID.fullmatch(txt) or (TS_ID.match(txt) if txt.startswith("QR_TS_SYS_") else None)
            if m:
                cur_head = m.group(0)
        elif tag == "tbl":
            tbl = tables[ti]; ti += 1
            if cur_head:
                d, e, tcs = parse_suite(tbl)
                suites.append((cur_head, d, e, tcs))
                cur_head = None
    return suites


# --------------------------------------------------------------------------- #
#  Safety lookup (RVTM / SRS)                                                  #
# --------------------------------------------------------------------------- #
_SAFETY_YES = {"safety related", "safety-related", "safety", "yes", "y", "true", "1", "sr"}
_SAFETY_NO  = {"not safety related", "not safety-related", "non safety related",
               "non-safety related", "not safety", "no", "n", "false", "0", "nsr"}


def _normalise_safety(v):
    s = str(v).strip().lower()
    if s in _SAFETY_YES:
        return "Safety Related"
    if s in _SAFETY_NO:
        return "Not Safety Related"
    return str(v).strip() if s else None


def load_safety(path, req_col=None, safety_col=None):
    """Return dict {SYS_id: 'Safety Related'|'Not Safety Related'|<raw>}."""
    rows = _read_table(path)
    if not rows:
        return {}
    header = rows[0]
    hlow = [str(h).strip().lower() for h in header]

    def find(name, fallback_pred):
        if name:
            for i, h in enumerate(hlow):
                if h == name.strip().lower():
                    return i
            sys.exit(f"ERROR: column '{name}' not found in {os.path.basename(path)}. "
                     f"Columns: {header}")
        return fallback_pred()

    def guess_req():
        best, score = None, -1
        for i in range(len(header)):
            hits = sum(1 for r in rows[1:] if r[i:i+1] and REQ_RE.fullmatch(str(r[i]).strip() or ""))
            if hits > score:
                best, score = i, hits
        return best

    def guess_safety():
        for i, h in enumerate(hlow):
            if "safety" in h:
                return i
        best, score = None, -1
        for i in range(len(header)):
            hits = sum(1 for r in rows[1:]
                       if str(r[i] if i < len(r) else "").strip().lower() in _SAFETY_YES | _SAFETY_NO)
            if hits > score:
                best, score = i, hits
        return best

    ci = find(req_col, guess_req)
    cs = find(safety_col, guess_safety)
    if ci is None or cs is None:
        sys.exit("ERROR: could not identify requirement / safety columns in the lookup. "
                 "Pass --req-col and --safety-col explicitly.")
    print(f"  Safety lookup: requirement column = '{header[ci]}', "
          f"safety column = '{header[cs]}'")
    out = {}
    for r in rows[1:]:
        if ci < len(r) and cs < len(r):
            rid = str(r[ci]).strip()
            if REQ_RE.fullmatch(rid):
                val = _normalise_safety(r[cs])
                if val:
                    out[rid] = val
    print(f"  Safety lookup: {len(out)} requirement classifications loaded.")
    return out


def _read_table(path):
    """Read xlsx/csv into a list of row-lists (first row = header)."""
    ext = os.path.splitext(path)[1].lower()
    if ext in (".xlsx", ".xlsm"):
        try:
            import openpyxl
        except ImportError:
            sys.exit("ERROR: reading .xlsx needs openpyxl (pip install openpyxl), "
                     "or export the lookup to .csv.")
        wb = openpyxl.load_workbook(path, data_only=True, read_only=True)

        def read_sheet(ws):
            rows, blank = [], 0
            for row in ws.iter_rows(values_only=True):
                cells = ["" if c is None else c for c in row]
                if all(str(c).strip() == "" for c in cells):
                    blank += 1
                    if blank > 50:
                        break
                    continue
                blank = 0
                rows.append(cells)
            return rows

        # pick the sheet whose cells contain the most SYS_#### requirement IDs
        best_rows, best_hits = None, -1
        for ws in wb.worksheets:
            data = read_sheet(ws)
            hits = sum(1 for r in data for c in r if REQ_RE.fullmatch(str(c).strip()))
            if hits > best_hits:
                best_rows, best_hits = data, hits
        return best_rows or []
    if ext == ".csv":
        import csv
        with open(path, newline="", encoding="utf-8-sig") as f:
            return [row for row in csv.reader(f)]
    sys.exit(f"ERROR: unsupported lookup type '{ext}'. Use .xlsx or .csv.")


# --------------------------------------------------------------------------- #
#  Document building                                                           #
# --------------------------------------------------------------------------- #
def _shade(cell, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def _add_borders(table):
    """Apply a full single-line grid to the table without needing 'Table Grid' style."""
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "0"); e.set(qn("w:color"), "808080")
        borders.append(e)
    tblPr.append(borders)


def _repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    th = OxmlElement("w:tblHeader"); th.set(qn("w:val"), "true")
    trPr.append(th)


def _set_widths(table, widths_twips):
    table.autofit = False
    for row in table.rows:
        for cell, w in zip(row.cells, widths_twips):
            cell.width = Pt(w / 20.0)  # twips -> points


def add_appendix_heading(doc, title):
    """Add an 'APPENDIX x.' heading without depending on a built-in style.
    Uses outline level 0 so it still appears in a Table of Contents / navigation."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    ol = OxmlElement("w:outlineLvl"); ol.set(qn("w:val"), "0"); pPr.append(ol)
    ks = OxmlElement("w:keepNext"); pPr.append(ks)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    # also tag it with the built-in Heading 1 style if the document defines one
    try:
        p.style = doc.styles["Heading 1"]
    except KeyError:
        pass


def add_matrix(doc, title, headers, data_rows, widths):
    add_appendix_heading(doc, title)
    tbl = doc.add_table(rows=1, cols=len(headers))
    _add_borders(tbl)
    hdr = tbl.rows[0]
    _repeat_header(hdr)
    for cell, text in zip(hdr.cells, headers):
        cell.text = ""
        p = cell.paragraphs[0]; run = p.add_run(text); run.bold = True
        run.font.size = Pt(9)
        _shade(cell, HEADER_FILL)
    for datum in data_rows:
        cells = tbl.add_row().cells
        for cell, text in zip(cells, datum):
            cell.text = ""
            r = cell.paragraphs[0].add_run("" if text is None else str(text))
            r.font.size = Pt(9)
    _set_widths(tbl, widths)
    doc.add_paragraph("")


def build(suites, safety, which, srs_rev, placeholder, has_rvtm=False):
    # aggregate
    tc_desc, tc_suites = {}, defaultdict(list)
    for sid, _d, _e, tcs in suites:
        for tid, td, _r in tcs:
            tc_desc.setdefault(tid, td)
            if sid not in tc_suites[tid]:
                tc_suites[tid].append(sid)

    appendices = []  # (letter, title, headers, rows, widths)
    excluded = []    # requirements screened out of Appendix A (not in RVTM)

    if "A" in which:
        rows, seen, seen_excl = [], set(), set()
        for sid, _d, _e, tcs in suites:
            for tid, _td, reqs in tcs:
                for rid in reqs:
                    # When an RVTM is supplied, requirements absent from it are
                    # highlighted and excluded from the matrix (not listed as TBC).
                    if has_rvtm and rid not in safety:
                        if rid not in seen_excl:
                            seen_excl.add(rid); excluded.append(rid)
                        continue
                    k = (rid, tid, sid)
                    if k not in seen:
                        seen.add(k)
                        rows.append((rid, safety.get(rid, placeholder), tid, sid))
        appendices.append((
            "A", "APPENDIX A. Requirement Traceability Matrix",
            ["Requirement ID", f"Safety related requirements from SRS revision {srs_rev}",
             "Test Case ID", "Test Suite ID"],
            rows, [1800, 3200, 2200, 2200]))

    if "B" in which:
        rows = []
        for sid, desc, execu, tcs in suites:
            tclist = []
            for tid, _td, _r in tcs:
                if tid not in tclist:
                    tclist.append(tid)
            rows.append((sid, desc, execu, ", ".join(tclist)))
        appendices.append((
            "B", "APPENDIX B. Test Suites Traceability Matrix",
            ["Test Suite", "Description", "Execution", "Associated Test Cases"],
            rows, [2400, 3600, 1200, 2200]))

    if "C" in which:
        rows = [(tid, tc_desc[tid], ", ".join(tc_suites[tid])) for tid in tc_desc]
        appendices.append((
            "C", "APPENDIX C. Test Cases Traceability Matrix",
            ["Test Case", "Description", "Associated Test Suites"],
            rows, [2200, 4200, 3000]))
    return appendices, excluded


def main():
    ap = argparse.ArgumentParser(
        description="Generate SPL traceability appendices (A/B/C) and append them to a copy of the export.",
        formatter_class=argparse.RawDescriptionHelpFormatter, epilog=__doc__)
    ap.add_argument("export", help="Test-suite export file (.doc or .docx)")
    ap.add_argument("--safety", help="RVTM/SRS lookup file (.xlsx or .csv) for the Safety-Related column")
    ap.add_argument("--req-col", help="Requirement-ID column name in the lookup (auto-detected if omitted)")
    ap.add_argument("--safety-col", help="Safety-class column name in the lookup (auto-detected if omitted)")
    ap.add_argument("--appendices", default="ABC", help="Subset of A B C to generate (default ABC)")
    ap.add_argument("--srs-rev", default="11.00 [5]", help="SRS revision text in Appendix A header")
    ap.add_argument("--placeholder", default="TBC", help="Value when safety class is unknown (default TBC)")
    ap.add_argument("-o", "--output", help="Output .docx path")
    args = ap.parse_args()

    which = [c for c in args.appendices.upper() if c in "ABC"]
    if not which:
        sys.exit("ERROR: --appendices must contain at least one of A, B, C.")

    print(f"Reading export: {args.export}")
    src_docx = ensure_docx(args.export)

    safety = load_safety(args.safety, args.req_col, args.safety_col) if args.safety else {}
    if not args.safety and "A" in which:
        print(f"  No --safety file given: Appendix A safety column set to '{args.placeholder}'.")

    print("Parsing test suites ...")
    suites = parse_export(src_docx)
    ntc = len({t[0] for _s, _d, _e, tcs in suites for t in tcs})
    nreq = len({r for _s, _d, _e, tcs in suites for _t, _td, rs in tcs for r in rs})
    print(f"  {len(suites)} test suites, {ntc} test cases, {nreq} requirements.")

    # work on a copy so the original export is never touched
    out = args.output or (os.path.splitext(args.export)[0] + "_with_appendices.docx")
    tmp = os.path.splitext(out)[0] + "_.tmp.docx"
    shutil.copyfile(src_docx, tmp)
    doc = docx.Document(tmp)
    doc.add_page_break()

    appendices, excluded = build(suites, safety, which, args.srs_rev,
                                 args.placeholder, has_rvtm=bool(args.safety))
    for letter, title, headers, rows, widths in appendices:
        print(f"  Appendix {letter}: {len(rows)} rows")
        add_matrix(doc, title, headers, rows, widths)

    if bool(args.safety) and "A" in which:
        if excluded:
            print(f"  Highlighted (not in RVTM, EXCLUDED from Appendix A): "
                  f"{len(excluded)} -> {', '.join(excluded)}")
        else:
            print("  RVTM coverage: every requirement is classified (nothing excluded).")

    doc.save(out)
    os.remove(tmp)
    print(f"\nDone -> {out}")


if __name__ == "__main__":
    main()
